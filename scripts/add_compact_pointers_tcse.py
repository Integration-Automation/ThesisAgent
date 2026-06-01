"""Add compact pointer sentences in TCSE_v2.3.docx's §3.2 + §6.2.

After the §1.3 末段 already carries a 61-char pointer to the thesis's
§3.7 / §6.4.5, the deck still only mentions the framework extensions
in ONE place. paper_inserts.md's intent was to mention the extensions
in three different anchor points (§1.3 / §3.2 / §6.2) so a reader who
skims any single section will still see the pointer to the thesis. The
full multi-paragraph version blew the 6-page budget, so this script
takes the middle path: add ONE compact sentence each to §3.2 + §6.2,
keeping each addition ≤ 90 chars to preserve the 6-page budget.

  §3.2 — append to the closing paragraph ("結合 RAG 規則檢索與 CoT
         多步推理之設計, ... 產出兼具可追溯性之程式碼審查結果。") a
         pointer sentence covering JudgeStep + author-feedback corpora,
         deferring quantitative evaluation to the thesis §3.5 / §3.7.

  §6.2 — append to the single §6.2 paragraph a pointer sentence covering
         the four inference backends + IDE-end MCP integration, deferring
         the cross-backend / feedback-corpus / deployment evaluations to
         the thesis §6.4.

Both appended sentences use 「，」 throughout (no 「；」 / `;`) per the
"Prose punctuation in additions" convention in CLAUDE.md.

Idempotent: each anchor pre-checks if the pointer is already present.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path("exports/TCSE_v2.3.docx")

# Pointer sentences — each ≤ ~90 chars; together ~150 chars budget.
S3_2_POINTER = (
    "本框架另含 JudgeStep（將模型裁決映射為 GitHub Review 事件之 APPROVE /"
    " REQUEST_CHANGES / COMMENT）與兩份作者反饋學習語料（dismissed /"
    " accepted）等設計貢獻，量化評估見學位論文 §3.5 與 §3.7。"
)

S6_2_POINTER = (
    "框架擴展方面，本研究隨附之開源實作另提供四類推論後端（本機 /"
    " FastAPI / OpenAI / Anthropic）與 IDE 端 MCP 整合介面，跨後端比較、"
    "作者反饋語料累積效益與部署層工程之驗證屬未來工作，見學位論文 §6.4。"
)


def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:50]!r}")


def _append_sentence_to_paragraph(paragraph, sentence: str) -> None:
    """Append text as a new run cloning rPr from an existing content run.
    Preserves font (eastAsia=標楷體, Times New Roman, sz=20)."""
    rpr_template = None
    for run in paragraph.runs:
        if (run.text or "").strip():
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                rpr_template = copy.deepcopy(rpr)
                break
    if rpr_template is None:
        for run in paragraph.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                rpr_template = copy.deepcopy(rpr)
                break
    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = sentence
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    paragraph._p.append(new_r)


def main() -> None:
    d = Document(SRC)
    before_chars = sum(len(p.text) for p in d.paragraphs)
    print(f"Opened: {SRC}  ({len(d.paragraphs)} paragraphs, {before_chars} chars)")

    # §3.2 — append to the "結合 RAG 規則檢索與 CoT 多步推理之設計" closing paragraph
    s32_idx = _find_para(d, "結合 RAG 規則檢索與 CoT 多步推理之設計")
    s32 = d.paragraphs[s32_idx]
    if "JudgeStep" in s32.text:
        print(f"  §3.2  [{s32_idx}]  skipped — pointer already present")
    else:
        _append_sentence_to_paragraph(s32, S3_2_POINTER)
        print(f"  §3.2  [{s32_idx}]  appended {len(S3_2_POINTER)} chars")

    # §6.2 — append to the single §6.2 content paragraph
    s62_idx = _find_para(d, "本研究仍有若干限制與發展方向")
    s62 = d.paragraphs[s62_idx]
    if "四類推論後端" in s62.text:
        print(f"  §6.2  [{s62_idx}]  skipped — pointer already present")
    else:
        _append_sentence_to_paragraph(s62, S6_2_POINTER)
        print(f"  §6.2  [{s62_idx}]  appended {len(S6_2_POINTER)} chars")

    d.save(SRC)
    after_chars = sum(len(p.text) for p in d.paragraphs)
    delta = after_chars - before_chars
    print(f"\nsaved: {SRC}  ({after_chars} chars, +{delta})")
    print(f"vs 6-page baseline (13841): {after_chars - 13841:+d}")
    if after_chars - 13841 > 250:
        print(f"  WARNING: over baseline by {after_chars - 13841} chars — may push past 6 pages")


if __name__ == "__main__":
    main()
