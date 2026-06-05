"""Fix three places that misframe CI/CD or version-control integration as
future work when the framework already has those capabilities.

User feedback (2026-06-01): "系統應用層面，未來可將本架構整合至版本控制
平台或 CI/CD 流程……這個我們已經有了，不會是未來研究工作". The exact
sentence is no longer in either docx (removed in earlier edits) but the
principle generalises: every place that hints at "integrating with CI/CD"
as future work needs the framing tightened, because §3.7.14 (a)–(p)
already documents the full CI matrix + GitHub Actions deployment, and
§3.6.2 documents the MCP / IDE integration.

Three edits — both files combined:

  論文 §6.3 (4) 部署面實證 — currently only acknowledges IDE
  integration. Rewrite to also cite §3.7.14 CI matrix as a done
  capability, so the limitation reads "both are built, neither is
  field-tested" rather than "only IDE is built".

  論文 §6.4.4 IDE 內審查觸發與生產級 ops 補強 — current first
  sentence reads "MCP 整合層使本框架可在 IDE 內直接觸發審查，後續可
  比較 IDE 觸發 vs CI 觸發兩種時機". A casual reader can parse "CI
  觸發" as "future capability". Rewrite to assert both triggers exist
  now (with cross-refs to §3.6.2 + §3.7.14), and frame ONLY the
  comparison study as future work.

  TCSE §6.2 末 — "跨後端比較、作者反饋語料累積效益與部署層工程之
  驗證屬未來工作" reads as "deployment-layer engineering verification is
  future work", ambiguous between "the engineering" and "its
  quantitative evaluation". Tighten to "之量化效益評估屬未來工作"
  matching the framing used in §3.7.14 / §1.5 item 6.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

THESIS = Path("exports/論文_v1.8.docx")
TCSE = Path("exports/TCSE_v2.3.docx")


def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:60]!r}")


def _content_rpr(paragraph):
    for run in paragraph.runs:
        if (run.text or "").strip():
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _rebuild_paragraph_text(paragraph, new_text: str) -> None:
    rpr_template = _content_rpr(paragraph)
    p_elem = paragraph._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)
    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    pPr = p_elem.find(qn("w:pPr"))  # NOSONAR OOXML WordprocessingML identifier (pPr/rPr/rFonts); lowercasing loses schema meaning
    if pPr is not None:
        pPr.addnext(new_r)
    else:
        p_elem.insert(0, new_r)


def _substring_replace(paragraph, old: str, new: str) -> bool:
    text = paragraph.text
    if old not in text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    _rebuild_paragraph_text(paragraph, text.replace(old, new, 1))
    return True


# ===========================================================================
# 論文 fix 1 — §6.3 (4) 部署面實證: add CI integration acknowledgment
# ===========================================================================

S6_3_4_OLD = (
    "(4) 部署面實證：本框架雖已整合至 IDE 環境完成可行性驗證，但尚未於"
    "真實開發團隊長期試行，缺乏對開發者採用意願、審查時間節省比例、"
    "誤報率與滿意度等實務指標之量化資料。"
)
S6_3_4_NEW = (
    "(4) 部署面實證：本框架已具備 IDE 端整合（§3.6.2 之 stdio MCP server）"
    "與 CI 端部署（§3.7.14 所述之 GitHub Actions matrix 分片、actions/cache"
    "串接 SQLite force-push 差分、per-PR 狀態快取與 idempotent 留言/check"
    " run 等部署層工程），但尚未於真實開發團隊長期試行，缺乏對開發者採用"
    "意願、審查時間節省比例、誤報率與滿意度等實務指標之量化資料。"
)


# ===========================================================================
# 論文 fix 2 — §6.4.4 first sentence: assert both triggers exist now
# ===========================================================================

S6_4_4_OLD = (
    "§3.6.2 所述之 MCP 整合層使本框架可在 IDE 內直接觸發審查，後續可比"
    "較 IDE 觸發（push 前）與 CI 觸發（push 後）兩種時機對開發者接受率"
    "與後續修正成本之差異。"
)
S6_4_4_NEW = (
    "本框架已具備兩種審查觸發路徑──§3.6.2 之 MCP 整合層支援 IDE 端 push"
    "前觸發、§3.7.14 之 GitHub Actions matrix 支援 CI 端 push 後觸發；"
    "後續可於累積實際 PR 流量後比較兩種時機對開發者接受率與後續修正"
    "成本之差異。"
)


# ===========================================================================
# TCSE fix 3 — §6.2 末: tighten "驗證" → "量化效益評估"
# ===========================================================================

TCSE_S6_2_OLD = "跨後端比較、作者反饋語料累積效益與部署層工程之驗證屬未來工作"
TCSE_S6_2_NEW = "跨後端比較、作者反饋語料累積效益與部署層工程之量化效益評估屬未來工作"


def update_thesis() -> None:
    d = Document(THESIS)
    before = sum(len(p.text) for p in d.paragraphs)
    print(f"=== 論文_v1.8.docx — {before} chars ===")

    # Fix 1: §6.3 (4)
    idx = _find_para(d, "(4) 部署面實證：本框架")
    para = d.paragraphs[idx]
    if "CI 端部署（§3.7.14" in para.text:
        print(f"  [{idx:3d}] §6.3 (4) skip — already updated")
    else:
        _rebuild_paragraph_text(para, S6_3_4_NEW)
        print(f"  [{idx:3d}] §6.3 (4): acknowledged CI integration is done")

    # Fix 2: §6.4.4
    idx = _find_para(d, "§3.6.2 所述之 MCP 整合層使本框架可在 IDE 內直接觸發審查")
    para = d.paragraphs[idx]
    if "本框架已具備兩種審查觸發路徑" in para.text:
        print(f"  [{idx:3d}] §6.4.4 skip — already updated")
    else:
        ok = _substring_replace(para, S6_4_4_OLD, S6_4_4_NEW)
        if not ok:
            raise SystemExit(f"§6.4.4 old text not found at [{idx}]")
        print(f"  [{idx:3d}] §6.4.4: clarified both triggers exist, only comparison is future")

    d.save(THESIS)
    after = sum(len(p.text) for p in d.paragraphs)
    print(f"  論文: {before} -> {after}  (Δ {after - before:+d})")


def update_tcse() -> None:
    d = Document(TCSE)
    before = sum(len(p.text) for p in d.paragraphs)
    print(f"\n=== TCSE_v2.3.docx — {before} chars ===")

    idx = _find_para(d, "本研究仍有若干限制與發展方向")
    para = d.paragraphs[idx]
    if TCSE_S6_2_NEW in para.text:
        print(f"  [{idx:3d}] §6.2 skip — already updated")
    else:
        ok = _substring_replace(para, TCSE_S6_2_OLD, TCSE_S6_2_NEW)
        if not ok:
            print(f"  [{idx:3d}] §6.2 skip — old substring not present (maybe TCSE has different wording)")
        else:
            print(f"  [{idx:3d}] §6.2: '驗證' → '量化效益評估'")

    d.save(TCSE)
    after = sum(len(p.text) for p in d.paragraphs)
    print(f"  TCSE: {before} -> {after}  (Δ {after - before:+d})")


def main() -> None:
    update_thesis()
    update_tcse()


if __name__ == "__main__":
    main()
