"""Revert TCSE_v2.3.docx's paper_inserts §1.1-§1.3 insertions and apply the
compact 1-sentence fallback specified in paper_inserts.md §3 audit checklist.

Context: applying the full §1.1 (3 paragraphs) + §1.2 (1 paragraph) +
§1.3 (1 paragraph) insertions pushed TCSE past the 6-page Word limit
(13,841 → 15,478 chars). paper_inserts.md's audit checklist gives an
explicit fallback for this case:

  「TCSE 短文版本可僅於 §1.3 之末句後追加一句『本研究隨附之開源框架
   另實作十三項研究級擴充機制，詳見學位論文 §3.7 與 §6.4.5，該等機制
   之量化評估均屬未來工作』以同步而不超頁。」

This script:
  1. Removes the 5 paragraphs the previous insert script added (matched
     by unique prefix so the revert is exact and idempotent).
  2. Appends the compact sentence to §1.3's last existing content
     paragraph (the one ending "...提升團隊協作效率並降低維護成本。").

The compact sentence uses 「，」 instead of 「；」 per the no-; convention
in CLAUDE.md "Content additions"; aside from that one substitution, the
wording is verbatim from paper_inserts.md.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path("exports/TCSE_v2.3.docx")

# Unique prefixes that identify each previously-inserted paragraph.
# Order doesn't matter for removal — we walk the body and remove any <w:p>
# whose visible text starts with any of these.
INSERTED_PREFIXES = [
    "在框架可擴展性方面，本研究將推論後端抽象為 Strategy 介面",
    "本研究隨附之開源框架另實作十三項研究級擴充機制（涵蓋 prompt-injection",
    "標示說明：上述段落僅描述機制設計",
    "框架另實作下列三項設計，於本研究實驗範圍外",
    "在框架擴展面，本研究隨附之開源實作目前已提供四類推論後端介面",
]

COMPACT_APPEND = (
    "本研究隨附之開源框架另實作十三項研究級擴充機制，詳見學位論文 §3.7 與"
    " §6.4.5，該等機制之量化評估均屬未來工作。"
)


def _remove_inserted_paragraphs(doc) -> int:
    """Remove every paragraph whose visible text starts with any inserted prefix.
    Returns the count removed."""
    body = doc.element.body
    removed = 0
    # Collect XML elements to remove, then remove (don't mutate while iterating).
    to_remove = []
    for p in doc.paragraphs:
        text = p.text.lstrip("　 \t")
        for prefix in INSERTED_PREFIXES:
            needle = prefix.lstrip("　 \t")
            if text.startswith(needle):
                to_remove.append(p._p)
                break
    for p_elem in to_remove:
        body.remove(p_elem)
        removed += 1
    return removed


def _append_sentence_to_paragraph(paragraph, sentence: str) -> None:
    """Append ``sentence`` to the end of ``paragraph``'s text without breaking
    its run-level formatting. We add a NEW run (cloning rPr from the existing
    content run) at the end, so the appended text inherits the same font."""
    # Find a representative rPr from an existing content run.
    rpr_template = None
    for run in paragraph.runs:
        if (run.text or "").strip():
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                rpr_template = copy.deepcopy(rpr)
                break
    if rpr_template is None:
        for run in paragraph.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                rpr_template = copy.deepcopy(rpr)
                break

    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = sentence
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    paragraph._p.append(new_r)


def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:50]!r}")


def main() -> None:
    d = Document(SRC)
    before_paras = len(d.paragraphs)
    before_chars = sum(len(p.text) for p in d.paragraphs)
    print(f"Opened: {SRC}  ({before_paras} paragraphs, {before_chars} chars)")

    removed = _remove_inserted_paragraphs(d)
    print(f"  removed {removed} previously-inserted paragraph(s)")

    # Idempotency: don't append the compact sentence twice.
    s13_last_idx = _find_para(d, "實驗結果顯示，本框架於 CRSCORE++ 指標上顯著優於基準")
    s13_last = d.paragraphs[s13_last_idx]
    if "本研究隨附之開源框架另實作十三項研究級擴充機制" in s13_last.text:
        print("  compact sentence already present — skipping append")
    else:
        _append_sentence_to_paragraph(s13_last, COMPACT_APPEND)
        print(f"  appended compact sentence to §1.3 last paragraph [{s13_last_idx}]"
              f"  (+{len(COMPACT_APPEND)} chars)")

    d.save(SRC)
    after_paras = len(d.paragraphs)
    after_chars = sum(len(p.text) for p in d.paragraphs)
    print(f"\nsaved: {SRC}  ({after_paras} paragraphs, {after_chars} chars)")
    print(f"net delta: {after_paras-before_paras:+d} paragraphs, "
          f"{after_chars-before_chars:+d} chars")
    print(f"baseline (pre-paper_inserts) was 13841 chars; now {after_chars}"
          f"  → {'WITHIN baseline' if after_chars <= 13841 else f'+{after_chars-13841} over baseline'}")


if __name__ == "__main__":
    main()
