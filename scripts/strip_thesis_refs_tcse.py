"""Remove all 「學位論文」 cross-references from TCSE_v2.3.docx.

Rationale: TCSE is a stand-alone conference paper, so it should not
forward-reference a separate "學位論文" — that framing weakens the
paper's self-contained framing and makes it look like a derivative
work. The framework-extension MENTIONS themselves stay (per the prior
"提到新東西的地方不夠多" feedback); only the cross-reference clauses
to §X of a thesis document get stripped or rewritten.

Two surviving 學位論文 mentions (audited 2026-05-31):

  §1.3 (para 26): "...十三項研究級擴充機制，詳見學位論文 §3.7 與
                   §6.4.5，該等機制之量化評估均屬未來工作。"
       →  drop the "詳見學位論文 §3.7 與 §6.4.5，" clause; the rest
          ("該等機制之量化評估均屬未來工作。") already conveys the
          deferral.

  §3.2 (para 51): "...等設計貢獻，量化評估見學位論文 §3.5 與 §3.7。"
       →  replace tail with "量化評估屬未來工作。" — same deferral
          framing, no cross-doc pointer.

§6.2 (para 90) already lost its 學位論文 reference in an earlier save
round-trip; no edit needed there.

Net delta: −27 chars (well under the 6-page budget headroom).
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path("exports/TCSE_v2.3.docx")


# (paragraph-prefix anchor, find_substring, replace_substring, label)
EDITS: list[tuple[str, str, str, str]] = [
    (
        "實驗結果顯示，本框架於 CRSCORE++ 指標上顯著優於基準",
        "本研究隨附之開源框架另實作十三項研究級擴充機制，詳見學位論文 §3.7 與 §6.4.5，該等機制之量化評估均屬未來工作。",
        "本研究隨附之開源框架另實作十三項研究級擴充機制，該等機制之量化評估均屬未來工作。",
        "§1.3 末: drop '詳見學位論文 §3.7 與 §6.4.5' clause",
    ),
    (
        "結合 RAG 規則檢索與 CoT 多步推理之設計",
        "等設計貢獻，量化評估見學位論文 §3.5 與 §3.7。",
        "等設計貢獻，量化評估屬未來工作。",
        "§3.2 末: rewrite '見學位論文 §3.5 與 §3.7' tail",
    ),
]


def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:50]!r}")


def _content_rpr(paragraph):
    for run in paragraph.runs:
        if (run.text or "").strip():
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _substring_replace(paragraph, old: str, new: str) -> bool:
    text = paragraph.text
    if old not in text:
        return False
    # In-run fast path.
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    # Cross-run fallback — merge to single run preserving rPr.
    rpr_template = _content_rpr(paragraph)
    new_text = text.replace(old, new, 1)
    p_elem = paragraph._p
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(new_r)
    else:
        p_elem.insert(0, new_r)
    return True


def main() -> None:
    d = Document(SRC)
    before = sum(len(p.text) for p in d.paragraphs)
    print(f"Opened: {SRC}  ({before} chars)")

    for anchor, old, new, label in EDITS:
        idx = _find_para(d, anchor)
        para = d.paragraphs[idx]
        old_len = len(para.text)
        if old not in para.text:
            if new in para.text:
                print(f"  [{idx:3d}] skip — already stripped ({label})")
            else:
                print(f"  [{idx:3d}] MISS — neither old nor new found ({label})")
            continue
        _substring_replace(para, old, new)
        delta = len(para.text) - old_len
        print(f"  [{idx:3d}] {old_len:4d} -> {len(para.text):4d}  ({delta:+d})  {label}")

    after = sum(len(p.text) for p in d.paragraphs)
    print(f"\nTOTAL: {before} -> {after}  ({after-before:+d})")
    print(f"vs 6-page baseline (13841): {after-13841:+d}")

    # Final audit: confirm zero 學位論文 mentions remain
    remaining = sum(1 for p in d.paragraphs if "學位論文" in p.text)
    print(f"\n學位論文 mentions remaining: {remaining}  "
          f"{'PASS' if remaining == 0 else 'FAIL'}")

    d.save(SRC)
    print(f"saved: {SRC}")


if __name__ == "__main__":
    main()
