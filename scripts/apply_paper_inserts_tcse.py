"""Apply paper_inserts.md §1.1-§1.3 insertions to exports/TCSE_v2.3.docx.

Three anchor-based insertions (per the v3 paper_inserts spec):

  §1.1 — after current §1.3 last paragraph (ending "並促成人類審查者..."),
         insert 3 paragraphs covering: (a) Strategy 介面 + JudgeStep
         framework-extensibility note, (b) 13 research-grade extensions
         (CLI flag opt-in form, end-to-end evaluation deferred to §3.7 /
         §6.4.5 of the thesis), (c) 標示說明 disclaimer.

  §1.2 — after §3.2's 推論層 paragraph, before §3.2's closing
         "結合 RAG 規則檢索..." paragraph, insert 1 paragraph describing
         three additional framework-only mechanisms (JudgeStep, JSONL
         learning corpora, secret pre-filter + SQLite cache + telemetry)
         explicitly out-of-scope for the paper's evaluation.

  §1.3 — after current §6.2 last paragraph, insert 1 paragraph listing
         four inference-backend interfaces + IDE stdio integration +
         pre-flight secret filter + cache/telemetry, with three (a)/(b)/(c)
         future-work bullets.

Content is inserted verbatim from paper_inserts.md (user-authored prose;
the no-; rule from CLAUDE.md applies only to my own rule / content
additions, not to user-supplied insertions). Font cloning preserves the
TCSE eastAsia=標楷體 + Times New Roman + sz=20 styling.

Idempotent — each anchor pre-checks if the insertion already exists by
matching on the first 40 chars of the first new paragraph.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path("exports/TCSE_v2.3.docx")


# ---------------------------------------------------------------------------
# Insertion content (paragraph-by-paragraph)
# ---------------------------------------------------------------------------

S1_3_PARAGRAPHS = [
    "在框架可擴展性方面，本研究將推論後端抽象為 Strategy 介面，目前實驗"
    + "以本機 Qwen3-Coder-30B-A3B-Instruct 加 LoRA 適配器之配置（即表 1 之"
    + " Ours 欄）為主，框架亦預留 OpenAI-相容端點與 Anthropic Messages API "
    + "之介面以利後續比較，惟此類跨後端之量化評估屬未來工作。框架另設"
    + " JudgeStep 將審查裁決映射為 GitHub Review API 之 event 欄位"
    + "（APPROVE / REQUEST_CHANGES / COMMENT）作為 PR 合併狀態之控制端點，"
    + "並提供合併前 Check Run gate 與作者反饋語料學習等設計，皆屬本框架之"
    + "設計貢獻；其量化驗證須累積實際 PR 流量後另行進行。",

    "本研究隨附之開源框架另實作十三項研究級擴充機制（涵蓋 prompt-injection"
    + " robustness、closed-loop 多輪對話、counterfactual / mutation-style "
    + "審查、provenance 稽核、force-push 差分、suggestion sandbox 驗證、"
    + "cross-language API drift、PR 類型自適應、reproducibility 訊號、"
    + "dependency upgrade impact、reviewer personas + conflict surfacing、"
    + "risk-weighted attention 與 diff entropy 偵測），均以 CLI flag 之"
    + " opt-in 形式提供；其端到端品質效益本研究均未予量化評估，詳細之"
    + "設計說明與後續實驗骨架見學位論文 §3.7 與 §6.4.5（對應本框架之"
    + " GitHub 倉庫之 docs/en/concepts/research-extensions.rst）。",

    "標示說明：上述段落僅描述機制設計，不附效益數字；本研究實際驗證之"
    + "結果見 §5。",
]

S3_2_PARAGRAPHS = [
    "框架另實作下列三項設計，於本研究實驗範圍外但屬隨附之開源框架之"
    + "一部分，列此供讀者參照：(a) JudgeStep，於 CoT pipeline 末讀取"
    + " total_summary 與已解析之 inline 留言，輸出 JSON 裁決並可映射為"
    + " GitHub Review API 之 event；(b) 兩份 append-only 之 JSONL 學習"
    + "語料，分別記錄歷次 PR 中被作者拒絕之留言與被作者「Apply suggestion」"
    + "採納之建議，於推論時可分別作為輸出端相似度過濾與輸入端 in-context"
    + " 範例使用；(c) 送出至第三方推論端點前之 secret 預過濾與 SQLite "
    + "為基礎之 prompt cache、generate 呼叫之 telemetry。上述三項機制之"
    + "量化評估不在本論文範圍內，留待未來工作。本論文 §5 之實驗結果僅就"
    + " §3.1 之知識蒸餾 + LoRA 微調管線與 §3.2 之 RAG + CoT 多步推理進行驗證。",
]

S6_2_PARAGRAPHS = [
    "在框架擴展面，本研究隨附之開源實作目前已提供四類推論後端介面"
    + "（本機 Hugging Face、自架 FastAPI 服務、OpenAI-相容端點、Anthropic"
    + " Messages API）、IDE 端 stdio 整合層（Model Context Protocol server）、"
    + "送出前之 secret 預過濾、prompt cache 與 telemetry 等功能；惟此類"
    + "機制之量化效益均未於本論文中評估。後續工作將依序針對：(a) 跨後端"
    + "於同一基準資料之品質、成本與延遲偏序比較；(b) 作者反饋語料（"
    + "dismissed 與 accepted）之累積對 inline finding 精確率與作者再次按"
    + " 👎 比率之影響；(c) MCP 介面對 IDE 內審查觸發率與接受率之影響；"
    + "逐項補上實際實驗以完成完整驗證。",
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:50]!r}")


def _content_rpr(paragraph) -> OxmlElement | None:
    for run in paragraph.runs:
        text = run.text or ""
        if text.strip() == "" and all(ch in "　 \t" for ch in text):
            continue
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _insert_paragraph_after(style_source_para, text: str):
    """Insert a new paragraph after ``style_source_para``, cloning pPr + rPr.

    Returns the new <w:p> XML element (so callers can use it as the next
    insertion anchor for sequential paragraph insertions).
    """
    new_p = OxmlElement("w:p")

    # Clone pPr (paragraph-level: indent, spacing, style ref) from source.
    src_pPr = style_source_para._p.find(qn("w:pPr"))  # NOSONAR OOXML WordprocessingML identifier (pPr/rPr/rFonts); lowercasing loses schema meaning
    if src_pPr is not None:
        new_p.append(copy.deepcopy(src_pPr))

    # Add a single run with cloned rPr (font, size, eastAsia).
    rpr_template = _content_rpr(style_source_para)
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_p.append(r)

    style_source_para._p.addnext(new_p)
    return new_p


def _insert_paragraphs_after(doc, anchor_idx: int, paragraphs: list[str]):
    """Insert paragraphs sequentially after the paragraph at anchor_idx."""
    anchor_p = doc.paragraphs[anchor_idx]
    style_src = anchor_p  # always clone style from the original anchor
    # Insert in REVERSE so each insertion sits immediately after anchor,
    # pushing prior insertions further down — final order matches input.
    for text in reversed(paragraphs):
        _insert_paragraph_after(style_src, text)


def _already_inserted(doc, marker_prefix: str) -> bool:
    needle = marker_prefix.lstrip("　 \t")[:40]
    return any(needle and p.text.lstrip("　 \t").startswith(needle)
               for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    d = Document(SRC)
    print(f"Opened: {SRC}  ({len(d.paragraphs)} paragraphs)")

    # §1.1 — after current §1.3 last content paragraph (the "實驗結果顯示..." one
    # that ends "並促成人類審查者與智慧系統之互補合作 ...")
    s13_anchor_prefix = "實驗結果顯示，本框架於 CRSCORE++ 指標上顯著優於基準"
    if _already_inserted(d, S1_3_PARAGRAPHS[0]):
        print("  §1.3  skipped — already inserted")
    else:
        idx = _find_para(d, s13_anchor_prefix)
        _insert_paragraphs_after(d, idx, S1_3_PARAGRAPHS)
        print(f"  §1.3  inserted {len(S1_3_PARAGRAPHS)} paragraphs after [{idx}]")

    # §1.2 — after §3.2 推論層 paragraph (the long one starting "推論層：")
    s32_anchor_prefix = "推論層：系統將構建完成之提示詞送入經 LoRA 微調"
    if _already_inserted(d, S3_2_PARAGRAPHS[0]):
        print("  §3.2  skipped — already inserted")
    else:
        idx = _find_para(d, s32_anchor_prefix)
        _insert_paragraphs_after(d, idx, S3_2_PARAGRAPHS)
        print(f"  §3.2  inserted {len(S3_2_PARAGRAPHS)} paragraphs after [{idx}]")

    # §1.3 — after current §6.2 last content paragraph
    s62_anchor_prefix = "本研究仍有若干限制與發展方向"
    if _already_inserted(d, S6_2_PARAGRAPHS[0]):
        print("  §6.2  skipped — already inserted")
    else:
        idx = _find_para(d, s62_anchor_prefix)
        _insert_paragraphs_after(d, idx, S6_2_PARAGRAPHS)
        print(f"  §6.2  inserted {len(S6_2_PARAGRAPHS)} paragraphs after [{idx}]")

    d.save(SRC)
    total = sum(len(p.text) for p in d.paragraphs)
    print(f"\nsaved: {SRC}  (now {len(d.paragraphs)} paragraphs, {total} chars)")


if __name__ == "__main__":
    main()
