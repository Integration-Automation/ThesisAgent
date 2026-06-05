"""Build exports/TCSE_v2.4.docx from TCSE_v2.3.docx.

Full revision driven by two sources (user request):
  1. .claude/agents/rules/paper_rule.md — the paper-writing handbook.
  2. https://code-review-framework.readthedocs.io (the prthinker framework docs)
     — authoritative, verified design facts for technical enrichment.

What changes (and why), all preserving v2.3's fonts (標楷體 CJK + Times New
Roman Latin) by cloning existing runs / setting the four rFonts slots:

  A. §2 相關研究 restructured into visible numbered subsections 2.1–2.4
     (paper_rule "numbered subsection headings must be visible" HARD rule;
     §2 previously had zero subsections) + a 2.3 literature comparison table
     (paper_rule §2.3) summarising already-cited works.
  B. Orphan reference [18] (Liu et al., data quality for code-review comment
     generation) cited in §3.1's data-cleaning sentence (paper_rule: every
     reference needs a matching in-text citation).
  C. §3.2 enriched with VERIFIED design facts from the prthinker docs:
     the six design patterns (Strategy/Factory/Template Method/Registry/
     Repository/DI), the RAG embedding model (Qwen3-Embedding-4B) + FAISS
     IndexFlatIP + default relevance threshold 0.7, the per-file
     inline_findings step, and a fuller JudgeStep (0–10 score + verdict +
     conservative PR-level aggregation) + dismissed (τ=0.85) / accepted
     (top-K=3) learned corpora. All are DESIGN-level facts; no experimental
     numbers are invented (no-fabrication HARD rule).
  D. §1.3 / §6.2: name the open-source framework (prthinker) once for
     traceability, and make §6.2's "十七項擴充機制" concrete with
     representative verified examples (all implemented w/ unit tests;
     quantitative evaluation = future work).
  E. Table renumber: the new §2.3 literature table becomes 表1, so the
     existing 表1/2/3 → 表2/3/4 with every in-text reference updated.
  F. Minor zh-tw lexical fix: 優化 → 最佳化 (paper_rule academic-vocabulary).

v2.3 is NOT overwritten — output is a separate exports/TCSE_v2.4.docx.
"""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path("exports/TCSE_v2.3.docx")
OUT = Path("exports/TCSE_v2.4.docx")

EA_FONT = "標楷體"
LAT_FONT = "Times New Roman"


# --------------------------------------------------------------------------
# Low-level run / paragraph helpers (preserve 標楷體 + Times New Roman)
# --------------------------------------------------------------------------
def _make_run(text: str, sz: int, bold: bool):
    """Build a <w:r> with the four-slot font set + size (half-points: 10pt→20)."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")  # NOSONAR OOXML WordprocessingML identifier (pPr/rPr/rFonts); lowercasing loses schema meaning
    rFonts = OxmlElement("w:rFonts")  # NOSONAR OOXML WordprocessingML identifier (pPr/rPr/rFonts); lowercasing loses schema meaning
    rFonts.set(qn("w:ascii"), LAT_FONT)
    rFonts.set(qn("w:hAnsi"), LAT_FONT)
    rFonts.set(qn("w:cs"), LAT_FONT)
    rFonts.set(qn("w:eastAsia"), EA_FONT)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    for tag in ("w:sz", "w:szCs"):
        e = OxmlElement(tag)
        e.set(qn("w:val"), str(sz))  # NOSONAR OOXML/Word style literal; a constant adds no value in this one-shot script
        rPr.append(e)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _strip_runs(p_elem):
    for tag in ("w:r", "w:hyperlink"):
        for r in p_elem.findall(qn(tag)):
            p_elem.remove(r)


def rebuild_paragraph(para, text: str, sz: int = 20, bold: bool = False):
    """Replace a python-docx Paragraph's runs with one styled run, keep pPr."""
    _strip_runs(para._p)
    para._p.append(_make_run(text, sz, bold))


def clone_blank_paragraph(template_para, text: str, sz: int = 20, bold: bool = False):
    """Deep-copy a template paragraph's <w:p> (keeps pPr: indent/spacing/numbering),
    strip its runs, and add one styled run carrying ``text``. Returns the <w:p>."""
    p_elem = copy.deepcopy(template_para._p)
    _strip_runs(p_elem)
    p_elem.append(_make_run(text, sz, bold))
    return p_elem


# --------------------------------------------------------------------------
# Load + capture stable references BEFORE any structural insertion
# --------------------------------------------------------------------------
shutil.copyfile(SRC, OUT)
doc = Document(str(OUT))
P = doc.paragraphs

tmpl_chapter = P[32]   # 三、LLM 程式碼審查架構設計  (chapter heading style)
tmpl_subhead = P[18]   # 1.1 研究背景                (subsection heading style)
tmpl_body = P[19]      # 1.1 body                    (indented body style)

# chapter-heading size (clone from an existing chapter heading run)
_chap_rpr = tmpl_chapter.runs[0]._r.find(qn("w:rPr"))
_chap_sz_el = _chap_rpr.find(qn("w:sz")) if _chap_rpr is not None else None
CHAP_SZ = int(_chap_sz_el.get(qn("w:val"))) if _chap_sz_el is not None else 22

para_28 = P[28]   # 1.4 body (+ glued 二、相關研究)
para_29 = P[29]   # §2 paragraph 1  → becomes 2.1 body
para_30 = P[30]   # §2 paragraph 2  → becomes 2.2 body

tmpl_table = doc.tables[1]   # 6×4, style "ab"  → clone for literature table


# ==========================================================================
# STEP 1 — in-place paragraph rewrites (no structural change; indices stable)
# ==========================================================================

# B + naming: §1.3 framework name + keep "十七項" claim
rebuild_paragraph(
    P[26],
    "　　實驗結果顯示，本框架於 CRSCORE++ 指標上顯著優於基準，可彌補傳統工具於高層次"
    "語義推理與設計合理性判斷上之不足，並促成人類審查者與智慧系統之互補合作，提升團隊"
    "協作效率並降低維護成本。本研究隨附之開源框架（prthinker）另實作十七項研究級擴充"
    "機制，該等機制之量化評估均屬未來工作。",
)

# B: §3.1 cite [18] at the data-quality sentence
rebuild_paragraph(
    P[34],
    "　　本研究將知識蒸餾（Knowledge Distillation, KD）與微調（Fine-tuning）流程建構為"
    "一完整管線。透過具備思維鏈（CoT）之提示詞向教師模型發出請求，教師模型為具高度推理"
    "能力之大型基礎模型，可同時產生最終答案與中間推理步驟，使輸出具備結構化且高品質之推理"
    "軌跡。其後將教師模型生成之內容整理為微調資料集（fine-tune data），經資料清洗、格式"
    "轉換（instruction-following 格式）、錯誤樣本過濾及人工或自動評估，以確保資料品質"
    "[18]。蒸餾核心階段以該批高品質資料對學生模型進行微調訓練，本研究採用之學生模型為 "
    "Qwen3-Coder-30B-A3B-Instruct，透過學習教師模型之輸出分佈與推理模式，使其於較小參數"
    "規模下逼近教師模型之效能，達成知識壓縮。微調採 QLoRA（Quantized Low-Rank "
    "Adaptation）技術，於低位元量化基礎上引入低秩適配器進行參數更新，降低顯示卡記憶體與"
    "計算成本，使 30B 等級模型亦可於有限硬體資源下高效訓練。",
)

# C: §3.2 intro — enrich the vague "軟體設計模式" with the six concrete patterns
rebuild_paragraph(
    P[40],
    "　　本研究提出之程式碼審查框架整合檢索增強生成（RAG）與思維鏈（CoT）多步推理機制，"
    "使大型語言模型於審查過程中得以參考領域知識並逐步推理，藉此提升審查結果之準確性與"
    "可解釋性。系統整體架構如圖 1 所示，依職責劃分為輸入、流程編排、檢索增強、提示模板、"
    "離線訓練與推論等六層：輸入層讀取待審查之原始程式碼並啟動整體框架，流程編排層將審查"
    "任務拆解為多個循序步驟並串接各步驟之推理脈絡，檢索增強層依程式碼語義檢索相關之審查"
    "規則並注入提示詞，提示模板層為各步驟構建對應之結構化提示詞，離線訓練層透過知識蒸餾"
    "與 LoRA 微調預先產出輕量化之審查模型，推論層執行各步驟之審查並彙整為最終報告。各層"
    "以明確介面溝通，並以對應之軟體設計模式約束實作邊界：推論後端以策略模式（Strategy）"
    "與工廠方法（create_backend）於本機 GPU、遠端 HTTP、OpenAI 與 Anthropic 後端間切換，"
    "審查步驟以樣板方法（Template Method）搭配註冊表（Registry，@register_step）擴充，"
    "檢索層以儲存庫模式（Repository）統一向量索引存取，整體流程則以相依性注入（Dependency "
    "Injection）組裝，藉此維持模組獨立並便於替換審查步驟、提示模板、檢索器或基座模型。",
)

# C: 檢索增強層 — add embedding model + FAISS index type + default threshold 0.7
rebuild_paragraph(
    P[43],
    "檢索增強層：同步啟動 RAG 系統，將原始碼送入 RAG 層取得相關規則。系統先以 "
    "Qwen3-Embedding-4B 嵌入模型將程式碼語義編碼，再由 FAISS（Facebook AI 釋出之向量"
    "索引庫，支援高速近似最近鄰搜尋）之內積索引（IndexFlatIP，配合 L2 正規化等同餘弦"
    "相似度）檢索可能相關之審查規則文件，並依預設相關性閾值 0.7（可經 --rag-threshold "
    "調整）判斷檢索結果之適用性：若存在足夠相關之規則便將其注入提示詞作為上下文補充，反之"
    "則直接進入後續審查階段，以避免低相關性資訊干擾模型判斷。",
    sz=20,
)
# keep List Paragraph indentation on P[43] (it was a List Paragraph; rebuild kept pPr)

# C: 推論層 — append the per-file inline_findings step (verified 6th step)
rebuild_paragraph(
    P[49],
    "推論層：系統將構建完成之提示詞送入經 LoRA 微調之 Qwen3-Coder-30B 模型，以執行該步驟"
    "之審查任務，每完成一步即判斷 CoT 流程是否結束：若尚未結束，當前結果回饋至下一輪提示詞"
    "之構建，使後續步驟延續前序推理脈絡，形成多步漸進之審查鏈，若所有步驟均已完成，則將"
    "累積之各步輸出注入 total_summary 提示詞，由模型彙整為結構化之最終審查報告。各提示模板"
    "對應產出一份獨立結果並即時寫入 .md：first_summary.md 記錄 PR 之變更概覽與影響範圍、"
    "first_code_review.md 提供初步審查所發現之可讀性與命名問題、linter.md 列出結構化之"
    "靜態分析訊息（含規則編號、嚴重度、行號與修正建議）、code_smell.md 條列設計層級之"
    "程式碼異味及其優先級，total_summary.md 彙整前述四份結果並輸出整體結論、綜合評估與"
    "最終合併決策建議。於逐檔（per-file）模式另提供 inline_findings 步驟，輸出 "
    "{行號, 嚴重度, 建議} 之 JSON，由 runner 轉換為 GitHub 行內審查留言。每一步推理產物"
    "皆可獨立檢視、追溯與驗證，確保長鏈推理中途失敗時前序成果不致遺失。",
)

# C: JudgeStep paragraph — enrich with verdict/score + corpora thresholds
rebuild_paragraph(
    P[51],
    "　　結合 RAG 規則檢索與 CoT 多步推理之設計，本框架可將領域知識動態納入審查，並由"
    "分步推理避免單次提示詞過載所致之資訊遺漏，產出兼具可追溯性之程式碼審查結果。本框架"
    "另含 JudgeStep：對每份審查輸出 0–10 分與 approve／request_changes／comment 之裁決"
    "並附理由，逐檔裁決再以保守規則彙整為 PR 級事件（任一檔要求變更即 REQUEST_CHANGES，"
    "全部通過才 APPROVE，其餘為 COMMENT），對映 GitHub Review API。系統並維護兩份由作者"
    "反饋累積之學習語料：dismissed 語料以餘弦相似度（門檻 τ=0.85）濾除作者曾否決之重複"
    "建議，accepted 語料則檢索作者曾採納之範例（預設 top-K=3）注入提示詞。上述設計貢獻"
    "均已於開源框架實作，其量化效益評估屬未來工作。",
)

# D: §6.2 future work — make the 17 extensions concrete with verified examples
rebuild_paragraph(
    P[90],
    "　　本研究仍有若干限制與發展方向。資料集方面，目前僅以 44 筆測試資料驗證，未來可擴展"
    "至涵蓋多程式語言、多框架與真實開源專案之大規模資料集，以提升泛化能力與評估可靠性。"
    "模型能力方面，可進一步探索不同模型架構或多模型協作（multi-agent）機制，使各模型分別"
    "負責特定審查面向，以提升整體分析品質。框架擴展方面，本研究隨附之開源實作提供四類推論"
    "後端（本機／FastAPI／OpenAI／Anthropic）與 IDE 端 MCP 整合介面，另實作十七項研究級"
    "擴充機制（如對抗式韌性測試、多輪對話審查、反事實／變異審查、審查者人格、儲存庫知識"
    "圖譜接地、由 dismissed／accepted 語料蒸餾規則之主動學習，以及以 --gate-on 設定之合併"
    "前 Check Run 關卡等），均已附單元測試；跨後端比較、作者反饋語料累積效益與上述擴充"
    "機制之量化評估屬未來工作。",
)

# E: table renumber in §5 body + captions (+1 because new §2.3 table is 表1)
rebuild_paragraph(
    P[71],
    "本研究以下列方式驗證 RQ1–RQ4：表 2 以 CRSCORE++ 之 comprehensiveness、conciseness、"
    "relevance 三項回答 RQ1（完整框架 vs CRSCORE++ 基準）與 RQ2（相同參數規模、僅多階段"
    "提示詞之效益），表 3 拆解完整方法、微調加單一提示詞、未微調加多階段提示詞三種設定，"
    "兩兩差距量化提示詞與微調之邊際貢獻以對應 RQ3，表 4 之人工評分對應 RQ4，補強可維護性"
    "與正確性等深層語義維度之評估可信度。",
)
rebuild_paragraph(P[72], "表2  CRSCORE++ 整體評估", sz=20, bold=True)
rebuild_paragraph(P[77], "表3 提示詞設計比較 (LLM 評分)", sz=20, bold=True)
rebuild_paragraph(P[80], "表4 提示詞設計比較 (人工評分)", sz=20, bold=True)
rebuild_paragraph(
    P[79],
    "RQ3：由表 3 之消融結果可見，多階段提示詞之貢獻達 +34 分，而 LoRA 微調之邊際貢獻僅約 "
    "+2 分，多階段提示詞為影響分數之主導因素。",
)
rebuild_paragraph(
    P[83],
    "RQ4：表 4 之人工評分結果與表 3 之 LLM 評分呈現一致之趨勢。於 Maintainability、"
    "Correctness、Multi-Review Coverage 與 conciseness 四個維度皆與 LLM 評分結果一致地"
    "支持完整方法之優勢，僅 Readability 略低，整體可作為對自動化評分之有效交叉驗證。",
)

# F: fix glued chapter heading on §1.4 body (remove trailing 二、相關研究)
rebuild_paragraph(
    para_28,
    "本論文章節安排如下：第二節回顧相關文獻，第三節說明所提出之 LLM 程式碼審查架構設計，"
    "第四節描述實驗設計，第五節呈現實驗結果與分析，第六節總結研究發現、限制與未來發展"
    "方向。",
)


# ==========================================================================
# STEP 2 — §2 restructure: subsection headings 2.1–2.4 + literature table
# ==========================================================================

# 2.1 body (rewrite para_29) — code review value + human limits
rebuild_paragraph(
    para_29,
    "　　程式碼審查（Code Review）為軟體開發流程中的關鍵環節，透過開發者之間的相互檢視與"
    "評估，確保程式碼之品質、可維護性與一致性，並促進知識分享與團隊協作。此過程有助於在"
    "早期發現邏輯錯誤、安全漏洞與效能問題，亦能建立統一的程式風格與開發文化。研究指出，"
    "程式碼審查可在整合測試前發現約 50% 至 70% 的缺陷[13]，且審查意見之品質會直接影響"
    "後續程式修正之成效[14]。然而，人工審查存在明顯限制：時間成本偏高，處理評論所需時間"
    "隨數量呈線性增加[4]，於大型且複雜之專案中更易形成瓶頸[20]，並對資深工程師造成額外"
    "負擔[5]。",
)

# 2.2 body (rewrite para_30) — LLM assistance + static analysis
rebuild_paragraph(
    para_30,
    "　　為緩解人工審查之負擔，研究社群逐漸導入大型語言模型（LLM）作為輔助工具，以提升"
    "審查效率並協助開發者理解陌生領域[2]，同時改善 Pull Request（開發者向主分支提出之"
    "程式碼合併申請）之審查延遲問題[9]。LLM 為基於 Transformer（以自注意力機制為核心之"
    "網路）架構之深度學習模型，透過預訓練與微調即可理解自然語言與程式語言，惟其仍面臨"
    "偏見與幻覺等問題，需搭配人工監督[15]，且在介面設計上以圖形化操作更具可用性[3]。"
    "另一方面，靜態分析工具（Static Analysis Tools, STAs）可在程式尚未執行前檢測語法"
    "錯誤、潛在漏洞與風格問題，並常整合於 CI/CD（自動建置、測試與發布之開發流程）中；"
    "惟傳統工具在規則僵化與語義理解方面存在限制[1]，其檢測能力有相當比例無法涵蓋 LLM "
    "所能辨識之問題[5]，且規則更新不易[12]、缺乏深層之建議能力[6]，凸顯 LLM 補足其不足"
    "之潛力。",
)

# new heading elements
h_chapter = clone_blank_paragraph(tmpl_chapter, "二、相關研究", sz=CHAP_SZ, bold=True)
h_21 = clone_blank_paragraph(tmpl_subhead, "2.1 程式碼審查與其限制", sz=20, bold=True)
h_22 = clone_blank_paragraph(tmpl_subhead, "2.2 LLM 輔助審查與靜態分析", sz=20, bold=True)
h_23 = clone_blank_paragraph(tmpl_subhead, "2.3 評估方法與模型最佳化技術", sz=20, bold=True)
h_24 = clone_blank_paragraph(tmpl_subhead, "2.4 研究缺口", sz=20, bold=True)

# 2.3 body — evaluation + RAG + PEFT/KD, ending by pointing at the table
b_23 = clone_blank_paragraph(
    tmpl_body,
    "　　在評估層面，LLM-as-a-Judge 方法透過 LLM 自動對生成結果進行評分或比較，其效果已"
    "被證實可與人類審查者相當[13]，並能彌補傳統指標（如 BLEU、ROUGE，基於字串重疊之自動"
    "評估指標）無法反映語義品質之不足[11]，相關研究亦提出多維評估指標以提升可靠性[7][10]。"
    "在知識注入方面，檢索增強生成（Retrieval-Augmented Generation, RAG）結合外部知識檢索"
    "與文本生成，可有效降低 LLM 幻覺並提升準確性，已應用於安全建議與開發輔助系統[19][8]。"
    "在模型最佳化方面，參數高效微調（Parameter-Efficient Fine-Tuning, PEFT）僅需調整少量"
    "參數即可降低訓練成本，其中 LoRA 透過低秩矩陣實現高效適配[16]，QLoRA 進一步結合量化"
    "技術使大型模型能在有限資源下完成微調[21][17]；知識蒸餾（Knowledge Distillation, KD）"
    "則將大型教師模型之能力轉移至小型學生模型，於降低計算成本之同時維持生成品質[22]。表 1 "
    "彙整代表性方法之策略與侷限，並對照本研究之定位。",
)

# 2.4 body — research gap (echoes §1, sets up §3)
b_24 = clone_blank_paragraph(
    tmpl_body,
    "　　綜觀上述，結合 LLM、靜態分析與模型最佳化技術有助於建立更高效且可擴展之程式碼審查"
    "機制，惟現有方法多採單一提示詞架構或受限於規則僵化，於語義理解、跨階段一致性與評估"
    "可信度三個面向仍存在缺口。本研究即藉由多階段 CoT 推理、RAG 知識注入與雙重 Judge "
    "評估，於上述三個面向補足此一缺口，其架構設計詳見第三節。",
)

# literature comparison table (clone tmpl_table = 6×4) → becomes 表1
lit_rows = [
    ("文獻", "方法", "優點", "侷限"),
    ("Jaoua 等人 [1]", "LLM 結合靜態分析器生成審查", "補足規則僵化、提升語義覆蓋",
     "單階段提示、設計合理性判斷有限"),
    ("Llama-Reviewer [16]", "PEFT／LoRA 微調自動審查", "少量參數逼近全量微調",
     "單一提示詞、未注入領域規範"),
    ("Yang 與 Chen [12]", "RAG 結合 LLM 自動審查", "動態注入規則、降低幻覺",
     "規則更新與跨階段一致性受限"),
    ("CRScore／CRScore++ [7][10]", "以程式碼主張與異味為基準之自動評估", "反映語義品質、可重現",
     "偏重評估面、未涵蓋多階段審查生成"),
    ("本研究（prthinker）", "多階段 CoT＋RAG 規則注入＋KD／QLoRA＋雙重 Judge",
     "跨階段一致、可解釋、可部署", "基準資料規模有限（屬未來工作）"),
]
lit_tbl = copy.deepcopy(tmpl_table._tbl)
# repopulate every cell, preserving each cell's own rPr (header vs body styling)
from docx.table import Table  # noqa: E402

lit_table = Table(lit_tbl, tmpl_table._parent)
for r_idx, row in enumerate(lit_table.rows):
    for c_idx, cell in enumerate(row.cells):
        text = lit_rows[r_idx][c_idx]
        first_p = cell.paragraphs[0]
        # drop any extra paragraphs in the cell
        for extra in cell.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
        runs = first_p.runs
        if runs:
            runs[0].text = text
            for extra_r in runs[1:]:
                extra_r._r.getparent().remove(extra_r._r)
        else:
            first_p._p.append(_make_run(text, 20, r_idx == 0))
        # left-align text columns for readability (template centred numeric cells)
        first_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
lit_caption = clone_blank_paragraph(tmpl_body, "表1 相關研究比較", sz=20, bold=True)
lit_caption_ppr = lit_caption.find(qn("w:pPr"))
# centre the caption like the other table captions are bold standalone lines
lit_cap_align = OxmlElement("w:jc")
lit_cap_align.set(qn("w:val"), "center")
if lit_caption_ppr is None:
    lit_caption_ppr = OxmlElement("w:pPr")
    lit_caption.insert(0, lit_caption_ppr)
lit_caption_ppr.append(lit_cap_align)

# --- wire elements into the tree ---
# before para_29:  二、相關研究 ; 2.1
para_29._p.addprevious(h_chapter)
para_29._p.addprevious(h_21)
# before para_30:  2.2
para_30._p.addprevious(h_22)
# after para_30 (reverse order):  2.3 head, 2.3 body, 表1 caption, lit table, 2.4 head, 2.4 body
para_30._p.addnext(b_24)
para_30._p.addnext(h_24)
para_30._p.addnext(lit_tbl)
para_30._p.addnext(lit_caption)
para_30._p.addnext(b_23)
para_30._p.addnext(h_23)

doc.save(str(OUT))
print(f"saved: {OUT}")
print(f"paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")
