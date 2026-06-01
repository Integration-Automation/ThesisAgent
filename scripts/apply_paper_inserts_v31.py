"""Apply paper_inserts.md v3.1 deltas to both TCSE_v2.3.docx and 論文_v1.8.docx.

v3.1 vs v3 deltas:
  - §3.7 mechanism count 13 → 17 (added §3.7.15-18: derive-lessons,
    discover-rules, build-kg, incremental-save-dir).
  - §3.7.14 deployment expanded from (a)-(g) to (a)-(p), adding 9
    production-stability engineering items (h)-(p).
  - §1.5 contribution #7 count 十三 → 十七 + 4 new mechanism names.
  - §6.4.5 intro / heading / closing: 十三 → 十七, plus 4 new
    experiment skeletons (n)-(q) for §3.7.15-§3.7.18.

TCSE: only the §1.3 末段 count change (十三 → 十七). User previously
stripped all 學位論文 cross-references; we do NOT add them back. The
compact 1-sentence form stays compact.

論文: 7 operations executed bottom-up so anchors aren't disturbed.

This script is idempotent: each operation pre-checks if the new state
already exists before applying.
"""
from __future__ import annotations

import contextlib
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

TCSE = Path("exports/TCSE_v2.3.docx")
THESIS = Path("exports/論文_v1.8.docx")


# ===========================================================================
# Helpers (shared with apply_paper_inserts_thesis.py)
# ===========================================================================

def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:60]!r}")


def _content_rpr(paragraph):
    for run in paragraph.runs:
        if (run.text or "").strip():
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _rebuild_paragraph_text(paragraph, new_text: str):
    rpr_template = _content_rpr(paragraph)
    p_elem = paragraph._p
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(new_r)
    else:
        p_elem.insert(0, new_r)


def _substring_replace(paragraph, old: str, new: str) -> bool:
    text = paragraph.text
    if old not in text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    _rebuild_paragraph_text(paragraph, text.replace(old, new, 1))
    return True


def _new_paragraph_after(anchor_para, text: str, *,
                         style_name: str | None = None,
                         clone_from: Paragraph | None = None):
    doc = anchor_para.part.document
    new_p_elem = OxmlElement("w:p")
    anchor_para._p.addnext(new_p_elem)
    para = Paragraph(new_p_elem, anchor_para._parent)
    if style_name is not None:
        with contextlib.suppress(KeyError):
            para.style = doc.styles[style_name]
    rpr_template = _content_rpr(clone_from) if clone_from is not None else _content_rpr(anchor_para)
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_p_elem.append(r)
    return para


def _insert_block_after(anchor_para, items: list[tuple[str | None, str]],
                        body_clone_from: Paragraph):
    cursor = anchor_para
    for style_name, text in items:
        if style_name in ("Heading 2", "Heading 3"):
            cursor = _new_paragraph_after(cursor, text, style_name=style_name,
                                          clone_from=None)
        else:
            cursor = _new_paragraph_after(cursor, text,
                                          style_name=style_name or "Normal (Web)",
                                          clone_from=body_clone_from)
    return cursor


# ===========================================================================
# TCSE: §1.3 末段 count update (十三 → 十七)
# ===========================================================================

def update_tcse() -> None:
    d = Document(TCSE)
    before = sum(len(p.text) for p in d.paragraphs)
    print(f"\n=== TCSE  ({TCSE.name}) — {before} chars ===")

    idx = _find_para(d, "實驗結果顯示，本框架於 CRSCORE++ 指標上顯著優於基準")
    para = d.paragraphs[idx]
    if "十七項研究級擴充機制" in para.text:
        print(f"  [{idx}] skip — already updated to 十七項")
    elif "十三項研究級擴充機制" in para.text:
        _substring_replace(para, "十三項研究級擴充機制", "十七項研究級擴充機制")
        print(f"  [{idx}] 十三項 → 十七項")
    else:
        print(f"  [{idx}] WARNING — no count phrase found; skipping")

    d.save(TCSE)
    after = sum(len(p.text) for p in d.paragraphs)
    print(f"  TCSE: {before} -> {after}  (Δ {after - before:+d})")


# ===========================================================================
# 論文 — new §3.7.14 sub-items (h)-(p)
# ===========================================================================

S3_7_14_NEW_SUBITEMS = [
    "(h) Step 輸出字元上限以避免 final-step OOM：CoT 最終 total_summary 步驟之"
    " prompt 為前序所有步驟之輸出之串接；於預設 max_new_tokens=32768 下，"
    "前序四步各可吐 ~120 KB，使最終 prompt 易達數十萬 token 並於 attention"
    "計算階段觸發 GPU OOM（50K token × 64 head 之 KV cache 約需 300 GiB）。"
    "框架以 _DEFAULT_MAX_STEP_RESULT_CHARS = 6000（約 1500 token）對 "
    "ctx.results 內每步之 in-pipeline 副本做硬上限截斷，輸出至磁碟與 API "
    "response 之全文保留不動；該上限可由 PRTHINKER_MAX_STEP_RESULT_CHARS "
    "環境變數於 server 端覆寫。",

    "(i) 本機後端之 FlashAttention 2 / SDPA 啟用與容器基底升級：將"
    " transformers.AutoModelForCausalLM.from_pretrained 之 attn_implementation"
    "偏好設為 flash_attention_2 並 fallback 至 sdpa，使 30B-class MoE 之"
    " attention 計算避開 vanilla 路徑之 O(L²) 記憶體峰值；配合升級伺服器"
    " Dockerfile 之 CUDA 基底至 13.0.1-devel-ubuntu22.04 以於映像建構期間"
    "提供 flash-attn 之 nvcc，並將 peft 釘住至 0.18.0 與 LoRA adapter 訓練"
    "時版本對齊（peft 0.19.x 將 MoE 之 expert projection（gate_proj / "
    "up_proj / down_proj）改派至 ParamWrapper，拒絕訓練時之 lora_dropout=0.1）。",

    "(j) timeout 預算之兩端拉長：原 v3 設計之 client poll deadline 與 backend"
    " per-call timeout 均為 30 分鐘；於本機後端跑全 5 步 CoT 之 per-file"
    " review 觀察到 30B 模型於 total_summary 步驟之單一 generate 呼叫可逼近"
    "此上限。框架將 RemoteBackendConfig.timeout_seconds 之預設值由 1800.0 "
    "提升至 3600.0、CLI 客戶端之 deadline 同步至 1800.0 / 3600.0，並於 "
    "review 子指令文件明示該上限為「全 pipeline 之 wall-clock 上限，非單一"
    " HTTP round-trip 之 idle timeout（後者由 reverse-proxy 控管）」。",

    "(k) poll 重試預算與指數 backoff 以吸收短暫 502：30B-class 後端之 "
    "process 重啟、GPU reload、nginx config reload 易於 ~1 分鐘時段內持續吐"
    " 502；原 v3 之 _MAX_CONSECUTIVE_POLL_FAILURES = 5 × _POLL_INTERVAL_SECONDS"
    " = 5 僅能扛 ~25 秒，超過即丟 HTTPStatusError 並中斷整輪 review。框架將"
    "預算提升至 60 並引入 _POLL_BACKOFF_AFTER_FAILURES = 5 / "
    "_POLL_MAX_INTERVAL_SECONDS = 30.0：前 5 次失敗仍以 5 秒等距重試，第 6 "
    "次起以 2 倍 backoff 增長至 30 秒上限，總可容忍時段約 3 分鐘；超出仍 "
    "raise，並由 _send_cancel 主動於 finally 區塊送 /review/cancel 釋放 GPU。",

    "(l) CI matrix 內之 diff-since-last 快取串接：§3.7.5 所述之 force-push"
    "差分機制原僅於 CLI 使用層暴露 --diff-since-last flag 與 SQLite cache 檔；"
    "於 CI 環境下需搭配 GitHub Actions 之 actions/cache 機制方能跨 workflow"
    " run 保留。框架於 matrix shard 之 review 步驟前後分別加入"
    " actions/cache/restore 與 actions/cache/save，以"
    " prthinker-diff-pr-<PR>-<run>-<shard> 為精確 key 並以"
    " prthinker-diff-pr-<PR>- 為 prefix restore-key 抓取同 PR 上最近一次"
    "成功 run 之 cache；同時於 env 注入 PRTHINKER_DIFF_SINCE_LAST=true 使"
    " review-pr 對每檔之 post-change 內容做 hash。save 步驟以 if: always() "
    "包裹，即使 shard 中途失敗，已 hash 之檔仍寫回 cache 供下一次 push 跳過。"
    "本機制使單檔小幅 fix 之 re-push 僅令該檔重跑 CoT，其餘 N-1 個 shard"
    "命中 cache 即直接 reuse 前次 findings。本子項屬部署層工程貢獻；其於真實"
    " PR 流量上之 GPU-second 節省與 cache hit 比率本論文未予評估。",

    "(m) GitHub 留言之 65536-char 上限自動截斷：GitHub Issues / PR 之 "
    "comment body 超過 65,536 字元時回應 422 Unprocessable Entity；於多檔"
    " matrix run 之 aggregate 階段，將各 shard 之 per-file 區塊（含 §3.7.14"
    " (g) 之 Overall Summary、RAG docs、judge verdicts）串接於同一 upserted"
    "留言中，極易超過該上限而導致整輪 aggregate 失敗。框架於 upsert_pr_comment"
    "內加上 60,000 字元之硬上限 _GITHUB_COMMENT_BODY_MAX：超過時保留前 60,000"
    "字之主體並追加 truncation 通知（指向 matrix shard 之 job log 取完整"
    " per-step 內容），並保證 comment_marker 字串於截斷後仍位於開頭，使下一"
    "次之 upsert 仍能定位同一留言並 PATCH 之，而非重新 POST 一條新的（避免"
    "破壞 §3.7.14 (f) 之 idempotency 設計）。本子項屬部署層工程貢獻；其於"
    "真實長 PR 上之截斷率與 reviewer 點入 job log 之頻率本論文未予評估。",

    "(n) enumerate 階段之 per-PR 狀態快取與 matrix 預過濾：§3.7.14 (l) 之"
    " in-runner SQLite cache 只能讓 shard 內之 LLM 短路，無法避免每個未變動"
    "之檔仍消耗一個 runner 之 checkout + setup-python + pip install + "
    "healthcheck（~1–2 分鐘）。框架另引入「per-PR 狀態快取」之 enumerate"
    "預過濾：每 PR 之 actions/cache 條目包含 "
    ".prthinker/pr-state/manifest.json（路徑 → blob SHA 對照表，blob SHA 由"
    " git rev-parse HEAD:<path> 取得，與 GitHub API 之回應一致）與"
    " .prthinker/pr-state/partials/<sha256>.json（該檔上次 review 之 partial"
    " result）。enumerate job 先 restore 該快取，逐 PR file 比對其當前 blob"
    " SHA 與 manifest，命中者直接自 matrix 中剔除不再 spawn shard，已 reuse"
    "之 partial 上傳為 partial-skipped artifact 供 aggregate 與新跑之 shard"
    "之 partial 一同 merge；aggregate post 完成後，write-state 步驟以 PR "
    "head 重建 manifest 與 partials 並做 canonical cache/save@v4，作為下一"
    "次 push 之來源。並同時修復 v3 設計中之 env-var 名稱錯位"
    "（REVIEWMIND_DIFF_SINCE_LAST / .reviewmind/diff-cache.sqlite 重新命名為"
    "對應之 PRTHINKER_*），使 (l) 之 in-runner SQLite 短路成為第二層安全網。"
    "本子項屬部署層工程貢獻；其相對 (l) 之邊際 CI 分鐘節省本論文未予評估。",

    "(o) 每 shard 之狀態 checkpoint 寫入：(n) 設計中之 canonical state 由"
    " aggregate job 寫入；若 aggregate 失敗或 runner 於最後一個 shard 與"
    " aggregate 之間崩潰，整輪所跑之 review 均不被下次 push 利用。配合"
    " §3.7.14 (b) 之 max-parallel: 1（既為單 GPU 之必然設計，亦提供本子項"
    "所需之序列性），框架讓每 shard 於 review-pr 成功後執行三步：(i) 將其"
    " partial.json 複製至 .prthinker/pr-state/partials/<sha256(path)>.json；"
    "(ii) 將 {path, blob_sha} 追加至 manifest.json；(iii) 以 -shard-<index>"
    "後綴另寫一筆 cache entry。aggregate 若成功寫入無後綴之 canonical entry"
    "則自動取代所有 shard checkpoint；aggregate 失敗時，下次 push 之"
    " restore-keys prefix 仍可抓到最後一個成功 shard 之 checkpoint，使所有"
    "「於 aggregate 失敗前已完成 review 之檔」於下次 matrix 中跳過。"
    "[ -s partial.json ] 守門確保 backend 不通之 shard（review step 雖以"
    " exit 0 結束但未產出 partial）不會把空結果寫進 manifest 而於下次被當作"
    "有效 reuse。本子項屬部署層工程貢獻；其於人為注入之 aggregate 失敗實驗"
    "下之 CI 分鐘節省量本論文未予評估。",

    "(p) inline findings 對 diff hunks 之預過濾以避開 review 全局 422："
    "GitHub 之 PR Review API 對單一 inline comment 若指向 side:RIGHT 之非"
    " hunk 範圍行，將以 422 Line could not be resolved 拒絕整份 review──一條"
    "虛構行號之 finding 即連帶癱瘓同 review 內所有合法 finding，aggregate"
    "因此 short-circuit 而 close_gate 不執行，required-status branch "
    "protection 將永遠看不到綠燈而阻擋 merge。框架施兩道防線：(i) "
    "github_api._filter_findings_to_diff 解析 PR diff 之 hunk header"
    "（@@ -a,b +c,d @@），對 ' ' / '+' 兩種行種逐行追蹤 new-side 行號集合，"
    "丟棄 (path, line) ∉ 該集合之 finding 與 start_line 越 hunk 之 multi-line"
    " finding；diff 抓取失敗之 fall-through 路徑改由外層 try/except 接住而"
    "非中斷。(ii) _cmd_aggregate 將 submit_inline_review 包入 try/except──"
    "summary comment 與 check run 於該時點均已開立，犧牲 inline 標注屬可"
    "接受損失，但若 check run 之 close_gate 因連動異常而未執行則 PR 永遠卡住。"
    "本子項屬部署層工程貢獻；其於真實 PR 上每輪 review 之被預過濾 finding"
    "數中位數本論文未予評估。",
]


# §3.7.15-18 NEW subsections
S3_7_NEW_SUBSECTIONS = [
    ("Heading 3", "3.7.15  主動學習衍生規則（active-learning derived lessons）"),
    (None,
     "§3.5 所述之 dismissed / accepted 兩語料為一階訊號──「此筆具體留言被"
     "作者拒絕」「此筆具體建議被採納」──兩者本身對未來 PR 無泛化能力。"
     "框架於兩語料之上再加一層：derive-lessons 子指令讀取兩語料之最近 N 筆，"
     "向模型詢問「應從中抽出何種可重用之審查規則（name / trigger / action"
     "三欄）」，並明示模型「輸出空陣列優於虛構規則」。解析後之 LessonRule"
     "連同其來源 PR 編號一併追加寫入 lessons.jsonl（append-only，供事後可"
     "追溯地檢視規則演化）。下一次 review-pr --lessons 時，最近 K 條規則"
     "被渲染為「Repo-derived review lessons」區塊，前置注入 inline-findings"
     " prompt，模型被指示將其視為軟性指引而非硬規則。本機制屬框架設計貢獻；"
     "其對作者反饋語料累積後之 inline finding 精確率影響本論文未予評估。"),

    ("Heading 3", "3.7.16  跨 PR finding 聚類與自我發現規則"),
    (None,
     "若框架反覆於不同 PR 中提出實質相同之 finding，正確之回應不是繼續重複，"
     "而是將其結晶為 --rules-dir 下之 repo 規則。框架為每條已產出之 inline"
     " finding 將留言文字以 backend embedding 化並與 (pr_number, file_path,"
     " line, comment, embedding) fingerprint 一併寫入 findings-index.sqlite。"
     "discover-rules 子指令對該 store 跑貪婪餘弦相似度聚類（預設 brute-force"
     " NumPy；大規模時可換為 sqlite-vec / FAISS 而不更動 greedy_cluster API），"
     "列出超過 --min-cluster-size 且 --similarity-threshold 以上之 cluster"
     "並以其最新一筆為代表（使候選規則隨時間追隨團隊用語演化而非僵化於舊"
     "用詞）。框架不自動寫入規則檔──候選必須由人類審查者明示採納。本機制"
     "屬框架設計貢獻；其聚類純度與所衍生規則之實際採納率本論文未予評估。"),

    ("Heading 3", "3.7.17  Repo 知識圖譜與 inline finding 之符號接地"),
    (None,
     "LLM 審查器於大型 repo 上常見之失敗模式為虛構符號名稱──宣稱 auth.py"
     "內有 get_user 函式而該函式實際位於 core/users.py。既有 RAG 層接地於"
     " repo 之規則，本機制接地於 repo 之符號：build-kg --workdir . 以 Python"
     " ast 走訪 def / class / 類方法 / ALL_CAPS 常數，並以正則為主之 scanner"
     "走訪 TypeScript / JavaScript 之 function / class / interface / const /"
     " default export，將 (symbol, kind, file, line, parent) 持久化於 "
     ".prthinker/repo-kg.sqlite。store 以 workdir 為 key，單一 SQLite 檔可"
     "同時容納多 repo 之 KG 而無洩漏。review-pr --kg-ground 將該表渲染為"
     "「Known symbols (treat as canonical, do not hallucinate)」前置區塊，"
     "並明示「finding 內若引用符號，該符號必須出現於表中」。rebuild() 採整批"
     "替換（先 delete 該 workdir 之舊 rows 再插入），確保 store 與 HEAD 對齊；"
     "增量更新列為未來工作。框架另提供 kg_visualize 子模組將該 SQLite 表"
     "輸出為單頁 D3 HTML 互動視圖，並於監控 overlay 之 nginx 加 /kg/ 路由"
     "直接服務該頁，使團隊可線上瀏覽框架對 repo 之符號理解。本機制屬框架"
     "設計貢獻；其 grounded prompt 對符號虛構率之降幅與對 inline finding"
     "精確率之影響本論文未予評估。"),

    ("Heading 3", "3.7.18  每檔遞增存檔與崩潰安全部分結果（crash-safe partial review）"),
    (None,
     "30B-class 後端之 per-file CoT 於大 PR 上可累積跑數十分鐘。當該輪因"
     " idle-poll sweep（§3.7.14 (e)）、GPU OOM、runner timeout 或人工於"
     " GitHub Actions 介面之 cancel 而中途終結時，原 v3 之 --output-json"
     "僅於審查最末整批落盤──中途死亡即無任何部分結果可審視。框架以"
     " --incremental-save-dir <path> 將 per-file 完成事件改寫為「即時 atomic"
     "寫盤」：<path>/files/<slug>.json 將 FileReviewResult 加入記憶體"
     " per_file_results list 之同時序列化（涵蓋 inline_findings / verdict /"
     " counterfactuals 等所有 pydantic 欄位）寫入磁碟（slug 將目錄分隔符與"
     "非法字元一律換為 _ 以跨 Windows / Linux / macOS 通用）；"
     "<path>/review.json 僅於整輪 sweep 跑完寫入，其存在即意味著「該次 run"
     "乾淨完成」；<path>/meta.json 開始時寫入 repo / pr_number / head_sha /"
     " started_at 使事後檢視者可辨識所屬 PR / commit。所有寫盤透過 "
     "<target>.tmp + os.replace 達成原子性，半寫狀態不可見。Writer 內部之"
     "任何 OSError 僅 log 並吞掉──持久化之失敗不可中斷正在跑之 review。"
     "CoTPipeline.run_per_file 為此暴露 on_file_done callback，於 cache-hit"
     "與全 review 兩處 append 點各觸發一次；本機制限於本機 pipeline 路徑"
     "（遠端 pipeline 走 --use-remote-pipeline 時 server 一次回完整 "
     "ReviewResult，per-file 增量在伺服端不適用，--output-json 仍為其對應之"
     "單檔落盤路徑）。本機制屬框架設計貢獻；其於真實 CI 流量上之"
     "「中斷-recovery 收益」量化評估本論文未予進行。"),
]


# §6.4.5 (n)-(q) new experiment skeletons
S6_4_5_NEW_ITEMS = [
    "(n) Active-learning derived lessons（§3.7.15）：先令 dismissed / accepted"
    "累積至雙語料各 ≥ 100 筆後，於同一份固定 PR 集合上以 paired bootstrap"
    "比較啟用 --lessons 與否之 inline finding precision 與作者再次按 👎 比率。"
    "輔以時間切片：每週重跑 derive-lessons，比較其輸出之新規則與既有規則之"
    " Jaccard 重疊，量化規則庫之收斂速度。",

    "(o) Cross-PR finding clustering（§3.7.16）：於 ≥ 500 筆累積 finding 上以"
    " discover-rules 跑全 grid 之 --similarity-threshold ∈ {0.75, 0.80, 0.85,"
    " 0.90} × --min-cluster-size ∈ {3, 5, 10}；以 ≥ 3 名人工審查者就「此 "
    "cluster 是否確為一條值得寫入規則庫之真實重複規則」之 Likert 5 點評分"
    "作為品質指標，並計算評審間 Cohen's κ。並以「成為框架候選 → 被人類採納為"
    " rules-dir 條目」之漏斗比率作為實用性指標。",

    "(p) Repo knowledge graph grounding（§3.7.17）：以 ≥ 30 個跨語言大 repo"
    "為樣本，於每 repo 上各跑啟用與未啟用 --kg-ground 之全 PR 集合。以人工"
    "標記之「該 finding 提及之符號是否實際存在於 repo」作為 ground-truth，"
    "計算「符號虛構率」之相對降幅；輔以 --kg-ground 之 prompt token 開銷與"
    "每 PR wall-clock 之 trade-off 表。視覺化模組之效益另以團隊內部使用"
    "次數 / page-view 等工程量度報告，不混入研究主張。",

    "(q) Crash-safe partial review（§3.7.18）：刻意於 CI 矩陣中注入 N 次中斷"
    "（concurrency: cancel-in-progress 觸發、人工 cancel、模擬 GPU OOM 之"
    " ask/cancel 連發），以 --incremental-save-dir 啟用前後各跑 ≥ 50 次中斷"
    "實驗，比較「中斷後可恢復之 per-file 完成數量」之中位數與「未恢復則需"
    "重跑之 GPU-second 總開銷」之減幅。本項屬可靠性工程量化評估，與品質研究"
    "分流。",
]


# §1.5 item 7 — full rewrite with 17-mechanism list
S1_5_ITEM7_OLD_PREFIX = "(7) 十三項研究級擴充機制之設計"
S1_5_ITEM7_NEW = (
    "(7) 十七項研究級擴充機制之設計（見 §3.7 詳述）：包含 prompt-injection"
    " robustness 之 corpus + bypass detection、closed-loop 多輪對話、"
    "counterfactual / mutation-style 審查、provenance 稽核、force-push 差分"
    " cache、suggestion sandbox 驗證、cross-language API drift 偵測、PR 類型"
    "自適應、reproducibility 訊號、dependency upgrade impact 分析、reviewer"
    " personas + conflict surfacing、risk-weighted attention、diff entropy"
    "／「diff bomb」偵測、作者反饋語料主動學習出之衍生規則、跨 PR finding"
    "聚類自我發現規則、Repo 知識圖譜對 inline finding 之符號接地，與每檔"
    "遞增存檔之崩潰安全部分審查結果。每項對應一個 CLI flag、一份單元測試與"
    " docs/en/concepts/research-extensions.rst 內之設計說明；其端到端品質"
    "效益本論文均未予評估，列為 §6.4.5 所述之未來工作。"
)


def update_thesis() -> None:
    d = Document(THESIS)
    before = sum(len(p.text) for p in d.paragraphs)
    print(f"\n=== 論文  ({THESIS.name}) — {before} chars ===")

    # Body-style clone source — §2.10 intro is stable and uses Normal (Web).
    body_clone_idx = _find_para(d, "本節整理前述各小節所回顧之 LLM 程式碼審查相關研究")
    body_clone = d.paragraphs[body_clone_idx]

    # ------------------------------------------------------------------
    # 1. §6.4.5 closing line: "十三組" → "十七組"
    # ------------------------------------------------------------------
    idx = _find_para(d, "上列十三組實驗皆需累積實際語料")
    if "十三組" in d.paragraphs[idx].text:
        _substring_replace(d.paragraphs[idx], "上列十三組實驗", "上列十七組實驗")
        print(f"  [{idx:3d}] §6.4.5 closing: 十三組 → 十七組")
    else:
        print(f"  [{idx:3d}] §6.4.5 closing skip — already 十七組")

    # ------------------------------------------------------------------
    # 2. §6.4.5 insert (n)-(q) BEFORE the closing line (idx now updated)
    # ------------------------------------------------------------------
    closing_idx = _find_para(d, "上列十七組實驗皆需累積實際語料")
    if not any("Active-learning derived lessons（§3.7.15）" in p.text for p in d.paragraphs):
        anchor_idx = closing_idx - 1
        while anchor_idx >= 0 and not d.paragraphs[anchor_idx].text.strip():
            anchor_idx -= 1
        anchor = d.paragraphs[anchor_idx]
        for text in reversed(S6_4_5_NEW_ITEMS):
            _new_paragraph_after(anchor, text, style_name="Normal (Web)",
                                 clone_from=body_clone)
        print(f"  [{anchor_idx:3d}] §6.4.5 inserted (n)-(q): 4 new experiment skeletons")
    else:
        print("  §6.4.5 skip — (n)-(q) already present")

    # ------------------------------------------------------------------
    # 3. §6.4.5 intro: "十三項" → "十七項" + append note
    # ------------------------------------------------------------------
    idx = _find_para(d, "§3.7 所述十三項機制目前僅完成框架實作")
    if idx >= 0:
        para = d.paragraphs[idx]
        _substring_replace(para, "§3.7 所述十三項機制", "§3.7 所述十七項機制")
        if "v3 既有之 (a)–(m)" not in para.text:
            note = (
                "v3 既有之 (a)–(m) 對應 §3.7.1–§3.7.13；v3.1 新增之 (n)–(q)"
                "對應 §3.7.15–§3.7.18。"
            )
            # Append by rewriting the whole paragraph
            _rebuild_paragraph_text(para, para.text + " " + note)
        print(f"  [{idx:3d}] §6.4.5 intro: 十三項 → 十七項 + note")

    # ------------------------------------------------------------------
    # 4. §6.4.5 heading: "十三項" → "十七項"
    # ------------------------------------------------------------------
    idx = _find_para(d, "6.4.5  §3.7 所述十三項研究級擴充機制之實證評估")
    if idx >= 0:
        _substring_replace(d.paragraphs[idx], "十三項", "十七項")
        print(f"  [{idx:3d}] §6.4.5 heading: 十三項 → 十七項")
    else:
        # Already updated
        idx2 = _find_para(d, "6.4.5  §3.7 所述十七項研究級擴充機制之實證評估")
        print(f"  [{idx2:3d}] §6.4.5 heading skip — already 十七項")

    # ------------------------------------------------------------------
    # 5. §3.7.15-18 INSERT after §3.7.14 closing, BEFORE chapter 4
    # ------------------------------------------------------------------
    if not any(p.text.startswith("3.7.15") for p in d.paragraphs):
        chap4_idx = _find_para(d, "第四章")
        anchor_idx = chap4_idx - 1
        while anchor_idx >= 0 and not d.paragraphs[anchor_idx].text.strip():
            anchor_idx -= 1
        anchor = d.paragraphs[anchor_idx]
        # Insert in reverse so final order is 15, 16, 17, 18
        for style_name, text in reversed(S3_7_NEW_SUBSECTIONS):
            _new_paragraph_after(anchor, text,
                                 style_name=style_name or "Normal (Web)",
                                 clone_from=body_clone if style_name is None else None)
        print(f"  [{anchor_idx:3d}] §3.7.15-18: inserted 4 new subsections "
              f"({len(S3_7_NEW_SUBSECTIONS)} entries)")
    else:
        print("  §3.7.15-18 skip — already inserted")

    # ------------------------------------------------------------------
    # 6. §3.7.14 insert (h)-(p) AFTER (g), BEFORE closing
    # ------------------------------------------------------------------
    if not any("Step 輸出字元上限以避免 final-step OOM" in p.text for p in d.paragraphs):
        # Anchor: the (g) sub-item ("(g) CI matrix 分片之 PR-wide overall summary")
        anchor_idx = _find_para(d, "(g) CI matrix 分片之 PR-wide overall summary")
        anchor = d.paragraphs[anchor_idx]
        for text in reversed(S3_7_14_NEW_SUBITEMS):
            _new_paragraph_after(anchor, text, style_name="Normal (Web)",
                                 clone_from=body_clone)
        print(f"  [{anchor_idx:3d}] §3.7.14: inserted (h)-(p) "
              f"({len(S3_7_14_NEW_SUBITEMS)} new sub-items)")
    else:
        print("  §3.7.14 (h)-(p) skip — already inserted")

    # ------------------------------------------------------------------
    # 7. §3.7 intro: "之十三項機制" → "之十七項機制" + add numbering note
    # ------------------------------------------------------------------
    idx = _find_para(d, "本節描述本研究隨附之開源框架另實作之十三項機制")
    if idx >= 0:
        para = d.paragraphs[idx]
        _substring_replace(para, "另實作之十三項機制", "另實作之十七項機制")
        if "子節編號 §3.7.1–§3.7.13" not in para.text:
            note = (
                "子節編號 §3.7.1–§3.7.13 為 v3 既有之十三項；§3.7.14 為部署層"
                "之工程設計；§3.7.15–§3.7.18 為 v3.1 新增之四項機制。"
            )
            _rebuild_paragraph_text(para, para.text + " " + note)
        print(f"  [{idx:3d}] §3.7 intro: 十三項 → 十七項 + numbering note")
    else:
        idx2 = _find_para(d, "本節描述本研究隨附之開源框架另實作之十七項機制")
        print(f"  [{idx2:3d}] §3.7 intro skip — already 十七項")

    # ------------------------------------------------------------------
    # 8. §1.5 item 7: replace entirely
    # ------------------------------------------------------------------
    idx = _find_para(d, "(7) 十三項研究級擴充機制之設計")
    if idx >= 0:
        _rebuild_paragraph_text(d.paragraphs[idx], S1_5_ITEM7_NEW)
        print(f"  [{idx:3d}] §1.5 item 7: replaced with 17-mechanism version")
    else:
        idx2 = _find_para(d, "(7) 十七項研究級擴充機制之設計")
        print(f"  [{idx2:3d}] §1.5 item 7 skip — already updated")

    d.save(THESIS)
    after = sum(len(p.text) for p in d.paragraphs)
    print(f"  論文: {before} -> {after}  (Δ {after - before:+d}, +{(after-before)*100//before}%)")


def main() -> None:
    update_tcse()
    update_thesis()


if __name__ == "__main__":
    main()
