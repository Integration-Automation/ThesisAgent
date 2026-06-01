"""Compress 4 existing TCSE paragraphs to free space for the new §3.2 + §6.2
pointer sentences while keeping the deck within the 6-page Word budget.

Constraint: §3.2 and §6.2 paragraphs both carry new pointer sentences
to the thesis §3.5 / §3.7 / §6.4 — those MUST NOT be touched here.
§2 literature review is also a locked zone (per compress_tcse_v22.py).
Trim targets are §1.3 contributions paragraph, §4.1 baseline-data
paragraph, §5.1 RQ-table-correspondence paragraph, and §6.1 conclusion
paragraph.

Per compress_tcse_v22.py "Strategy: preserve ALL substantive content;
only strip AI tells, filler, stacked-template phrasing". Numbers,
hyperparameters, table references, and design rationale all kept.

Expected total trim: ~305 chars (bring 14,162 → ~13,857, only ~16 chars
over the original 13,841 baseline → should fit 6 pages cleanly).
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path("exports/TCSE_v2.3.docx")


# Each entry: (anchor_prefix, old_substring_or_full_paragraph, new_substring_or_full_paragraph, mode)
# mode = "substring" → in-paragraph replace; "paragraph" → full rewrite.
EDITS: list[tuple[str, str, str, str]] = [

    # §1.3 contribution paragraph — trim the "於錯誤檢測... 產出兼具一致性..." clause.
    # The follow-up "系統設計採多階段..." sentence already covers the審查結果
    # qualities (可追溯、可解釋), so the redundant qualifier in the opening
    # sentence is the cleanest target.
    (
        "針對上述動機，本研究建立一套結合 LLM",
        "於錯誤檢測、風格規範建議、最佳實踐推廣及語義理解等任務上，產出"
        "兼具一致性、可解釋性與可維護性之審查結果",
        "於錯誤檢測、風格規範與語義理解等任務上產出一致、可解釋之審查結果",
        "substring",
    ),

    # §4.1 baseline-data paragraph — rewrite full paragraph.
    # Cuts: (a) the "確保整體方法於不同程式碼情境下均具備穩定性、準確性與
    # 可重現性" tail of the opening sentence (filler — the rest of the
    # paragraph IS that justification); (b) "造成之風格偏差與分佈集中問題"
    # → "造成之風格偏差" (the two are the same issue); (c) the duplicate
    # multi-dim-problem list "（命名不一致、重複邏輯、過度耦合、例外處理
    # 缺失）" (already listed under bug / 效能 / 安全 / code smell);
    # (d) closing "用於分析不同模型於相同輸入下之輸出差異與能力保留程度"
    # which is implied by "對照樣本".
    (
        "本研究之實驗設計涵蓋基準資料建構",
        "本研究之實驗設計涵蓋基準資料建構、模型評估方法、人工審查機制與"
        "系統配置，確保整體方法於不同程式碼情境下均具備穩定性、準確性與"
        "可重現性。基準資料以 GPT-5 與 Copilot 共同生成 44 筆測試資料"
        "並透過人工驗證機制確保品質與正確性，避免單一資料來源造成之"
        "風格偏差與分佈集中問題，使測試資料涵蓋語法結構、命名習慣、"
        "設計模式及錯誤型態（潛在 bug、效能瓶頸、安全風險、code smell，"
        "功能正常但結構不良之程式特徵）等多樣情境，貼近真實開發場景。"
        "44 筆規模於覆蓋多維問題（命名不一致、重複邏輯、過度耦合、例外"
        "處理缺失）與控制人工標註與驗證成本間取得平衡，亦可作為教師模型"
        "與學生模型之對照樣本，用於分析不同模型於相同輸入下之輸出差異"
        "與能力保留程度。",
        "本研究之實驗設計涵蓋基準資料建構、評估方法、人工審查與系統配置。"
        "基準資料以 GPT-5 與 Copilot 共同生成 44 筆測試資料並經人工驗證"
        "以確保品質，避免單一資料來源造成之風格偏差，使測試資料涵蓋語法"
        "結構、命名習慣、設計模式及錯誤型態（潛在 bug、效能瓶頸、安全"
        "風險、code smell，功能正常但結構不良之程式特徵）等多樣情境。"
        "44 筆規模在覆蓋多維問題與控制人工標註成本間取得平衡，亦可作為"
        "教師模型與學生模型之對照樣本。",
        "paragraph",
    ),

    # §5.1 RQ-table correspondence paragraph — rewrite. Cuts the opening
    # "為釐清四個研究問題與實驗結果之對應關係" filler ("以下列方式驗證
    # RQ1-RQ4" implies the same), and "以排除模型容量差異後分離出提示詞之
    # 單獨貢獻" + "判斷分數提升之主因" — these are restatements of what
    # 表 2's two-pair-diff structure already says.
    (
        "為釐清四個研究問題與實驗結果之對應關係",
        "為釐清四個研究問題與實驗結果之對應關係，本研究依下列方式驗證："
        "表 1 以 CRSCORE++ 之 comprehensiveness、conciseness、relevance "
        "三項自動化指標回答 RQ1（本研究完整框架之多階段提示詞加 LoRA "
        "微調 vs CRSCORE++ 基準）與 RQ2（相同參數規模、僅多階段提示詞之"
        "效益），以排除模型容量差異後分離出提示詞之單獨貢獻，表 2 拆解"
        "組合完整方法、微調加單一提示詞、未微調加多階段提示詞三種設定，"
        "兩兩差距分別量化提示詞與微調之邊際貢獻以對應 RQ3，判斷分數提升"
        "之主因，表 3 為相同三組設定之人工評分，作為對 LLM-as-a-Judge "
        "之交叉驗證以對應 RQ4，補強可維護性、正確性等深層語義維度之"
        "評估可信度。",
        "本研究以下列方式驗證 RQ1–RQ4：表 1 以 CRSCORE++ 之"
        " comprehensiveness、conciseness、relevance 三項回答 RQ1（完整框架"
        " vs CRSCORE++ 基準）與 RQ2（相同參數規模、僅多階段提示詞之效益），"
        "表 2 拆解完整方法、微調加單一提示詞、未微調加多階段提示詞三種"
        "設定，兩兩差距量化提示詞與微調之邊際貢獻以對應 RQ3，表 3 之人工"
        "評分對應 RQ4，補強可維護性與正確性等深層語義維度之評估可信度。",
        "paragraph",
    ),

    # §6.1 conclusion paragraph — rewrite. Heavy duplication with §1.3
    # (the multi-stage pipeline + RAG + KD/QLoRA combo is already laid
    # out in the intro); the conclusion can compress to "we did X, it
    # works under resource constraints" without re-listing the 5 steps.
    # Also removes the prose ; per the no-; convention.
    (
        "本研究針對傳統人工程式碼審查於效率、成本與一致性上之限制",
        "本研究針對傳統人工程式碼審查於效率、成本與一致性上之限制，提出"
        "一套結合大型語言模型與思維鏈推理之多階段自動化審查框架，將流程"
        "拆分為摘要生成、初步審查、靜態分析、程式碼異味偵測與最終報告"
        "整合等階段，並搭配結構化提示詞設計，降低單一提示之負載，提升"
        "分析深度、穩定性與可解釋性；同時結合檢索增強生成以引入專案規範"
        "與外部知識，減少模型幻覺並提升結果可靠度。模型訓練與部署面採"
        "知識蒸餾與參數高效微調（QLoRA），使系統於有限資源下仍維持具"
        "實用價值之審查能力，兼顧效能與成本效益。",
        "本研究針對傳統人工程式碼審查於效率、成本與一致性之限制，提出"
        "結合大型語言模型與思維鏈推理之多階段自動化審查框架，並搭配"
        "結構化提示詞設計與檢索增強生成以引入專案規範、降低模型幻覺。"
        "模型訓練面採知識蒸餾與 QLoRA 微調，使系統於有限資源下維持實用"
        "之審查能力，兼顧效能與成本效益。",
        "paragraph",
    ),
]


def _find_para(doc, prefix: str) -> int:
    needle = prefix.lstrip("　 \t")
    for i, p in enumerate(doc.paragraphs):
        if p.text.lstrip("　 \t").startswith(needle):
            return i
    raise SystemExit(f"anchor not found: {prefix[:50]!r}")


def _content_rpr(paragraph):
    for run in paragraph.runs:
        if (run.text or "").strip():
            rpr = run._r.find(qn("w:rPr"))
            if rpr is not None:
                return copy.deepcopy(rpr)
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def _rebuild_paragraph_text(paragraph, new_text: str) -> None:
    rpr_template = _content_rpr(paragraph)
    p_elem = paragraph._p
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)
    new_r = OxmlElement("w:r")
    if rpr_template is not None:
        new_r.append(rpr_template)
    t = OxmlElement("w:t")
    t.text = new_text
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(new_r)
    else:
        p_elem.insert(0, new_r)


def _substring_replace(paragraph, old: str, new: str) -> bool:
    """In-place substring replace within a paragraph's runs.

    First try in-run; fall back to paragraph rebuild (merging runs) if
    the substring crosses runs.
    """
    text = paragraph.text
    if old not in text:
        return False
    # In-run fast path
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    # Cross-run fallback — merge into single run, preserves font.
    _rebuild_paragraph_text(paragraph, text.replace(old, new, 1))
    return True


def main() -> None:
    d = Document(SRC)
    before = sum(len(p.text) for p in d.paragraphs)
    print(f"Opened: {SRC}  ({before} chars)")

    total_saved = 0
    for anchor, old, new, mode in EDITS:
        idx = _find_para(d, anchor)
        para = d.paragraphs[idx]
        old_len = len(para.text)
        if mode == "paragraph":
            if para.text.strip() == new.strip():
                print(f"  [{idx:3d}] skip — already compressed")
                continue
            _rebuild_paragraph_text(para, new)
        else:  # substring
            if new in para.text and old not in para.text:
                print(f"  [{idx:3d}] skip — already compressed")
                continue
            ok = _substring_replace(para, old, new)
            if not ok:
                raise SystemExit(f"substring not found in para [{idx}]: {old[:30]!r}")
        new_len = len(para.text)
        saved = old_len - new_len
        total_saved += saved
        print(f"  [{idx:3d}] {old_len:4d} -> {new_len:4d}  saved {saved:+d}")

    after = sum(len(p.text) for p in d.paragraphs)
    print(f"\nTOTAL saved: {total_saved} chars  ({before} -> {after})")
    print(f"vs 6-page baseline (13841): {after - 13841:+d}")

    d.save(SRC)
    print(f"saved: {SRC}")


if __name__ == "__main__":
    main()
